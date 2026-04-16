"""Pipeline orchestrator.

Ties together the scraper, differ, and updater modules into a single end-to-end
workflow:

    sources.json + guide.docx/md
            │
            ▼
    ┌───────────────┐
    │  1. Scrape     │  For each approved URL, fetch the latest page text
    │     (scraper)  │  and compare against the previous snapshot.
    └───────┬───────┘
            │ changed sources
            ▼
    ┌───────────────┐
    │  2. Diff       │  For every source that changed, produce a structured
    │     (differ)   │  summary of additions / removals.
    └───────┬───────┘
            │ combined diff
            ▼
    ┌───────────────┐
    │  3. Update     │  Ask an LLM to rewrite only the affected sections
    │     (updater)  │  of the Sponsor Guide.
    └───────┬───────┘
            │
            ▼
    output/sponsor_guide_updated.md  +  .docx
"""

import argparse
import json
import logging
import os
import re
from pathlib import Path
from typing import Optional

import mammoth
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from dotenv import load_dotenv

import cite
import differ
import scraper
import updater
from src.utils.logging_utils import configure_rotating_file_logging
from src.utils.source_policy import assert_public_sources

try:
    import markdown as _markdown_lib
    from fpdf import FPDF as _FPDF
    _PDF_AVAILABLE = True
except Exception:
    _PDF_AVAILABLE = False

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# I/O helpers
# ---------------------------------------------------------------------------

def load_sources(config_path: str) -> list[dict]:
    """Load approved source URLs from a JSON config file.

    Expected format::

        [
          {"name": "NIH_R15", "url": "https://...", "data_class": "public"},
          ...
        ]

    Args:
        config_path: Path to the JSON file.

    Returns:
        List of validated source dicts with ``name``, ``url``, ``sections``,
        and ``data_class``.
    """
    with open(config_path, "r", encoding="utf-8") as fh:
        sources = json.load(fh)

    if not isinstance(sources, list):
        raise ValueError("sources.json must contain a JSON array")

    for idx, src in enumerate(sources):
        if not isinstance(src, dict):
            raise ValueError(f"Source at index {idx} must be an object")

        name = str(src.get("name", "")).strip()
        url = str(src.get("url", "")).strip()
        if not name or not url:
            raise ValueError(f"Source at index {idx} must include non-empty 'name' and 'url'")

        sections = src.get("sections", [])
        if sections is None:
            sections = []
        if not isinstance(sections, list):
            raise ValueError(f"Source at index {idx} must use a list for 'sections'")

        data_class = str(src.get("data_class", "")).strip().lower()
        if data_class != "public":
            raise ValueError(
                f"Source '{name}' at index {idx} must set data_class to 'public' before it can be used."
            )

        src["name"] = name
        src["url"] = url
        src["sections"] = sections
        src["data_class"] = "public"

    return sources


def read_guide(path: str) -> str:
    """Read a Sponsor Guide and return its content as markdown.

    Supports ``.docx`` (converted via *mammoth*) and ``.md`` / ``.txt``
    (read verbatim).

    Args:
        path: File path to the guide.

    Returns:
        Markdown string.

    Raises:
        ValueError: If the file extension is not supported.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        with open(path, "rb") as fh:
            result = mammoth.convert_to_markdown(fh)
            if result.messages:
                for msg in result.messages:
                    logger.warning("event=mammoth_message detail=%s", msg)
            return result.value
    if ext in (".md", ".txt"):
        with open(path, "r", encoding="utf-8") as fh:
            return fh.read()
    raise ValueError(
        f"Unsupported guide format '{ext}'. Use .docx, .md, or .txt."
    )


def write_guide_md(path: str, content: str) -> None:
    """Write the updated markdown guide to *path*."""
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(content)


# ---------------------------------------------------------------------------
# Markdown → .docx converter (basic)
# ---------------------------------------------------------------------------

def _md_to_docx(md_text: str, path: str) -> None:
    """Convert markdown text to a ``.docx`` file.

    Handles headings, paragraphs, bold / italic, bullet lists, and simple
    pipe-delimited tables — enough for a typical Sponsor Guide.
    """
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # --- headings ---
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
            i += 1
            continue

        # --- table ---
        if "|" in line and i + 1 < len(lines) and re.match(r"^[\s|:-]+$", lines[i + 1]):
            headers = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for ci, h in enumerate(headers):
                p = table.rows[0].cells[ci].paragraphs[0]
                p.text = ""
                _add_inline_formatting(p, h)
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci < len(headers):
                        p = table.rows[ri + 1].cells[ci].paragraphs[0]
                        p.text = ""
                        _add_inline_formatting(p, cell)
            continue

        # --- bullet list (including indented sub-bullets) ---
        m = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, m.group(1))
            i += 1
            continue

        # --- numbered list ---
        m = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, m.group(1))
            i += 1
            continue

        # --- blank line ---
        if not line.strip():
            i += 1
            continue

        # --- regular paragraph ---
        p = doc.add_paragraph()
        _add_inline_formatting(p, line)
        i += 1

    doc.save(path)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    r_pr.append(u)
    r_pr.append(color)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_emphasis_runs(paragraph, text: str) -> None:
    """Parse **bold** and *italic* in plain text segments."""
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            run.bold = True
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_inline_formatting(paragraph, text: str) -> None:
    """Parse links + emphasis and add styled runs."""
    # Handles both [text](url) and [[1]](url) style citation links.
    link_pat = re.compile(r"\[([^\]]*(?:\[[^\]]*\])?[^\]]*)\]\((https?://[^)\s]+)\)")
    pos = 0
    for m in link_pat.finditer(text):
        if m.start() > pos:
            _add_emphasis_runs(paragraph, text[pos : m.start()])
        _add_hyperlink(paragraph, m.group(1), m.group(2))
        pos = m.end()
    if pos < len(text):
        _add_emphasis_runs(paragraph, text[pos:])


# ---------------------------------------------------------------------------
# Snapshot reader
# ---------------------------------------------------------------------------

def _read_snapshot(name: str, data_dir: str) -> str:
    """Return the previously saved text snapshot, or ``""`` on first run."""
    path = os.path.join(data_dir, f"{name}_latest.txt")
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as fh:
            return fh.read()
    return ""


# ---------------------------------------------------------------------------
# Markdown → PDF converter
# ---------------------------------------------------------------------------

# Characters outside fpdf2's default Latin-1 core fonts, mapped to safe equivalents.
_UNICODE_REPLACEMENTS = {
    "\u2018": "'",   # left single quotation mark
    "\u2019": "'",   # right single quotation mark
    "\u201a": ",",   # single low-9 quotation mark
    "\u201c": '"',   # left double quotation mark
    "\u201d": '"',   # right double quotation mark
    "\u201e": '"',   # double low-9 quotation mark
    "\u2013": "-",   # en dash
    "\u2014": "--",  # em dash
    "\u2015": "--",  # horizontal bar
    "\u2026": "...", # horizontal ellipsis
    "\u00a0": " ",   # non-breaking space
    "\u00b7": "*",   # middle dot
    "\u2022": "*",   # bullet
    "\u2023": "*",   # triangular bullet
    "\u2032": "'",   # prime
    "\u2033": '"',   # double prime
}


def _sanitize_for_pdf(text: str) -> str:
    """Replace Unicode characters unsupported by fpdf2's Latin-1 core fonts."""
    import unicodedata
    for char, replacement in _UNICODE_REPLACEMENTS.items():
        text = text.replace(char, replacement)
    # Normalize and drop any remaining non-Latin-1 characters.
    result = []
    for char in text:
        if ord(char) <= 255:
            result.append(char)
        else:
            normalized = unicodedata.normalize("NFKD", char)
            ascii_equiv = normalized.encode("ascii", "ignore").decode("ascii")
            result.append(ascii_equiv if ascii_equiv else "?")
    return "".join(result)


_LIST_ITEM_RE = re.compile(r"^(\s*)([-*+]|\d+\.)\s")


def _ensure_blank_before_lists(md_text: str) -> str:
    """Insert a blank line before the first list item in any run of list items
    that immediately follows a non-blank, non-list line.

    The Python ``markdown`` library needs this gap to recognise bullets/numbers
    as ``<ul>``/``<ol>`` instead of folding them into the preceding paragraph.
    """
    lines = md_text.split("\n")
    result: list[str] = []
    for i, line in enumerate(lines):
        if _LIST_ITEM_RE.match(line) and i > 0:
            prev = result[-1] if result else ""
            if prev.strip() and not _LIST_ITEM_RE.match(prev):
                result.append("")
        result.append(line)
    return "\n".join(result)


def _md_to_pdf(md_text: str, path: str) -> None:
    """Convert markdown to a styled PDF using fpdf2 (pure Python, zero system dependencies).

    Raises:
        ImportError: If markdown or fpdf2 are not installed.
    """
    if not _PDF_AVAILABLE:
        raise ImportError(
            "PDF export requires 'markdown' and 'fpdf2'. "
            "Run: pip3 install markdown fpdf2"
        )

    # Sanitize Unicode before converting so fpdf2's Latin-1 fonts don't choke.
    safe_md = _sanitize_for_pdf(md_text)

    # The Python markdown library requires a blank line before the first list
    # item; without it, bullets are swallowed into the preceding <p> tag.
    safe_md = _ensure_blank_before_lists(safe_md)

    # Convert markdown → HTML, then feed to fpdf2's HTML renderer.
    body_html = _markdown_lib.markdown(
        safe_md,
        extensions=["tables", "fenced_code", "sane_lists"],
    )

    # fpdf2 needs explicit width on <table> tags for proper column sizing,
    # and left-align all cells explicitly (fpdf2 defaults to center).
    body_html = body_html.replace("<table>", '<table width="100%">')
    body_html = body_html.replace("<th>", '<th align="left">')
    body_html = body_html.replace("<td>", '<td align="left">')

    full_html = f"<html><body>{body_html}</body></html>"

    pdf = _FPDF(orientation="P", unit="pt", format="Letter")
    pdf.set_margins(left=72, top=72, right=72)
    pdf.set_auto_page_break(auto=True, margin=72)
    pdf.add_page()
    pdf.write_html(full_html)
    pdf.output(path)


# ---------------------------------------------------------------------------
# Pipeline
# ---------------------------------------------------------------------------

def run_pipeline(
    sources_config: str,
    guide_path: str,
    output_dir: str = "output",
    model_name: str = updater.DEFAULT_MODEL,
    state_file: Optional[str] = None,
    data_dir: Optional[str] = None,
    with_citations: bool = True,
    citation_model: Optional[str] = None,
    refresh_citations: bool = False,
    refresh_citations_only: bool = False,
) -> bool:
    """Execute the full scrape → diff → update pipeline.

    Args:
        sources_config: Path to the JSON file listing approved source URLs.
        guide_path: Path to the current Sponsor Guide (``.docx`` or ``.md``).
        output_dir: Directory where updated guide files are written.
        model_name: LLM model identifier forwarded to the updater.

    Returns:
        ``True`` if the guide was updated, ``False`` if no changes were found.
    """
    # -- 1. Load inputs -------------------------------------------------------
    sources_root = os.path.dirname(os.path.abspath(sources_config)) or "."
    if state_file is None:
        state_file = os.path.join(sources_root, "state.json")
    if data_dir is None:
        data_dir = os.path.join(sources_root, "data")

    logger.info("step=1 action=load_inputs status=start")
    sources = load_sources(sources_config)
    guide_md = read_guide(guide_path)
    assert_public_sources(sources, context="pipeline")
    logger.info(
        "step=1 action=load_inputs status=done sources=%d guide_path=%s state_file=%s data_dir=%s",
        len(sources),
        guide_path,
        state_file,
        data_dir,
    )

    # -- 2. Scrape & diff -----------------------------------------------------
    all_diffs: list[tuple[str, list[str], str]] = []
    if refresh_citations_only:
        refresh_citations = True
        logger.info("step=2 action=scrape_diff status=skipped reason=refresh_citations_only")
    else:
        logger.info("step=2 action=scrape_diff status=start")
        for src in sources:
            name, url = src["name"], src["url"]
            sections = src.get("sections", [])
            old_text = _read_snapshot(name, data_dir)

            try:
                changed = scraper.check_for_updates(
                    url,
                    name,
                    state_file=state_file,
                    data_dir=data_dir,
                )
            except Exception as exc:
                logger.warning("step=2 source=%s status=scrape_failed error=%s", name, exc)
                continue

            if changed:
                new_text = _read_snapshot(name, data_dir)
                diff = differ.extract_changes(old_text, new_text)

                if not sections:
                    try:
                        sections = updater.classify_sections(new_text, guide_md)
                        if sections:
                            logger.info(
                                "step=2 source=%s status=sections_autodetected sections=%s",
                                name,
                                ",".join(sections),
                            )
                    except Exception:
                        pass

                all_diffs.append((name, sections, diff))
                logger.info("step=2 source=%s status=changed", name)
            else:
                logger.info("step=2 source=%s status=unchanged", name)

    if not all_diffs and not refresh_citations:
        logger.info("result=up_to_date reason=no_source_changes")
        return False

    updated_md = guide_md
    did_llm_update = False
    if all_diffs:
        # -- 3. Update via LLM ------------------------------------------------
        logger.info("step=3 action=llm_update status=start diffs=%d model=%s", len(all_diffs), model_name)

        diff_blocks: list[str] = []
        for name, sections, diff in all_diffs:
            header = f"## Source: {name}"
            if sections:
                header += (
                    "\nRelevant guide sections: "
                    + ", ".join(f'"{s}"' for s in sections)
                )
            diff_blocks.append(f"{header}\n\n{diff}")
        combined_diff = "\n\n".join(diff_blocks)

        try:
            assert_public_sources(sources, context="pipeline update step")
            updated_md = updater.update_guide(guide_md, combined_diff, model_name)
            did_llm_update = True
        except EnvironmentError as exc:
            logger.error("step=3 action=llm_update status=error type=config error=%s", exc)
            return False
        except Exception as exc:
            logger.error("step=3 action=llm_update status=error type=llm_call error=%s", exc)
            return False
    else:
        logger.info("step=3 action=llm_update status=skipped reason=no_diffs")

    # -- 4. Optional citation pass --------------------------------------------
    evidence: list[dict] = []
    if with_citations:
        assert_public_sources(sources, context="pipeline citation step")
        logger.info("step=4 action=citation_pass status=start")
        snapshot_map: dict[str, str] = {}
        for src in sources:
            name = src["name"]
            txt = _read_snapshot(name, data_dir)
            if not txt and not refresh_citations_only:
                try:
                    txt = scraper.fetch_and_clean_text(src["url"])
                except Exception:
                    txt = ""
            snapshot_map[name] = txt
        try:
            cited_md, evidence = cite.add_citations(
                updated_md,
                sources=sources,
                snapshots_by_name=snapshot_map,
                model_name=citation_model or model_name,
            )
            if evidence:
                updated_md = cited_md
                logger.info("step=4 action=citation_pass status=done claims=%d", len(evidence))
            else:
                logger.warning("step=4 action=citation_pass status=no_validated_citations")
        except Exception as exc:
            logger.warning(
                "step=4 action=citation_pass status=failed error=%s fallback=continue_without_citations",
                exc,
            )

    # -- 5. Write outputs ------------------------------------------------------
    step_label = "[5/5]" if with_citations else "[4/4]"
    logger.info("step=%s action=save_outputs status=start", step_label.strip("[]"))
    os.makedirs(output_dir, exist_ok=True)

    md_path = os.path.join(output_dir, "sponsor_guide_updated.md")
    write_guide_md(md_path, updated_md)
    logger.info("step=5 artifact=markdown status=saved path=%s", md_path)

    docx_path = os.path.join(output_dir, "sponsor_guide_updated.docx")
    try:
        _md_to_docx(updated_md, docx_path)
        logger.info("step=5 artifact=docx status=saved path=%s", docx_path)
    except Exception as exc:
        logger.warning("step=5 artifact=docx status=failed error=%s", exc)

    pdf_path = os.path.join(output_dir, "sponsor_guide_updated.pdf")
    try:
        _md_to_pdf(updated_md, pdf_path)
        logger.info("step=5 artifact=pdf status=saved path=%s", pdf_path)
    except ImportError as exc:
        logger.warning("step=5 artifact=pdf status=skipped reason=dependency_missing error=%s", exc)
    except Exception as exc:
        logger.warning("step=5 artifact=pdf status=failed error=%s", exc)

    if with_citations and evidence:
        evidence_path = os.path.join(output_dir, "sponsor_guide_evidence.json")
        with open(evidence_path, "w", encoding="utf-8") as fh:
            json.dump(evidence, fh, indent=2)
        logger.info("step=5 artifact=evidence status=saved path=%s", evidence_path)

    if did_llm_update:
        logger.info("result=updated changed_sources=%s", ",".join(n for n, _, _ in all_diffs))
    else:
        logger.info("result=citations_refreshed")
    return True


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    """Entry point when invoked from the command line."""
    configure_rotating_file_logging(log_file=Path("logs") / "pipeline.log")
    load_dotenv(Path(__file__).resolve().parent / ".env")

    parser = argparse.ArgumentParser(
        description="Sponsor Guide update pipeline — scrape → diff → LLM update",
    )
    parser.add_argument(
        "guide",
        help="Path to the current Sponsor Guide (.docx or .md)",
    )
    parser.add_argument(
        "--sources",
        default="sources.json",
        help="Path to the approved-sources JSON config (default: sources.json)",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Output directory (default: <sources_dir>/output/)",
    )
    parser.add_argument(
        "--state",
        default=None,
        help="Path to state.json (default: <sources_dir>/state.json)",
    )
    parser.add_argument(
        "--data-dir",
        default=None,
        help="Directory for text snapshots (default: <sources_dir>/data/)",
    )
    parser.add_argument(
        "--model",
        default=updater.DEFAULT_MODEL,
        help=f"LLM model name (default: {updater.DEFAULT_MODEL})",
    )
    parser.add_argument(
        "--with-citations",
        dest="with_citations",
        action="store_true",
        help=argparse.SUPPRESS,
    )
    parser.add_argument(
        "--no-citations",
        dest="with_citations",
        action="store_false",
        help="Disable citation pass (enabled by default).",
    )
    parser.add_argument(
        "--citation-model",
        default=None,
        help="Model used for citation mapping (default: same as --model).",
    )
    parser.add_argument(
        "--refresh-citations",
        action="store_true",
        help=(
            "Run citation pass even when no source diffs are detected "
            "(guide text is left unchanged)."
        ),
    )
    parser.add_argument(
        "--refresh-citations-only",
        action="store_true",
        help=(
            "Skip scrape/diff and only regenerate citations/evidence for the current guide "
            "using existing snapshots in the program data folder."
        ),
    )
    parser.set_defaults(with_citations=True)
    args = parser.parse_args()

    output_dir = args.output
    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(os.path.abspath(args.sources)), "output")

    run_pipeline(
        args.sources,
        args.guide,
        output_dir,
        args.model,
        state_file=args.state,
        data_dir=args.data_dir,
        with_citations=args.with_citations,
        citation_model=args.citation_model,
        refresh_citations=args.refresh_citations,
        refresh_citations_only=args.refresh_citations_only,
    )


if __name__ == "__main__":
    main()
