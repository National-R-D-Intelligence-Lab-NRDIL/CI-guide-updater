"""Markdown to DOCX converter.

Handles headings, paragraphs, bold/italic, bullet lists, numbered lists,
blockquotes, pipe-delimited tables, hyperlinks, and weekly-update red
highlight spans.
"""

import re
from typing import Optional

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

WEEKLY_HIGHLIGHT_COLOR = "C1121F"
_WEEKLY_HIGHLIGHT_SPAN_RE = re.compile(
    r"<span\s+style=[\"']color:\s*#?c1121f;?[\"']>(.*?)</span>",
    flags=re.IGNORECASE | re.DOTALL,
)


def _is_html_comment_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("<!--") and stripped.endswith("-->")


def _add_hyperlink(paragraph, text: str, url: str, color_hex: str = "0000FF") -> None:
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    r_pr.append(u)
    r_pr.append(color)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_emphasis_runs(paragraph, text: str, color_hex: Optional[str] = None) -> None:
    """Parse **bold** and *italic* in plain text segments."""
    def _style_run(run) -> None:
        if color_hex:
            run.font.color.rgb = RGBColor.from_string(color_hex)

    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            _style_run(run)
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
            _style_run(run)
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            run.bold = True
            _style_run(run)
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.italic = True
            _style_run(run)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _style_run(run)


def _add_inline_formatting(paragraph, text: str, color_hex: Optional[str] = None) -> None:
    """Parse links, emphasis, and weekly-update red spans."""
    pos = 0
    for m in _WEEKLY_HIGHLIGHT_SPAN_RE.finditer(text):
        if m.start() > pos:
            _add_inline_formatting(paragraph, text[pos : m.start()], color_hex=color_hex)
        _add_inline_formatting(paragraph, m.group(1), color_hex=WEEKLY_HIGHLIGHT_COLOR)
        pos = m.end()
    if pos:
        if pos < len(text):
            _add_inline_formatting(paragraph, text[pos:], color_hex=color_hex)
        return

    link_pat = re.compile(r"\[([^\]]*(?:\[[^\]]*\])?[^\]]*)\]\((https?://[^)\s]+)\)")
    pos = 0
    for m in link_pat.finditer(text):
        if m.start() > pos:
            _add_emphasis_runs(paragraph, text[pos : m.start()], color_hex=color_hex)
        _add_hyperlink(paragraph, m.group(1), m.group(2), color_hex=color_hex or "0000FF")
        pos = m.end()
    if pos < len(text):
        _add_emphasis_runs(paragraph, text[pos:], color_hex=color_hex)


def md_to_docx(md_text: str, path: str) -> None:
    """Convert markdown text to a .docx file.

    Handles headings, paragraphs, bold/italic, bullet lists, and simple
    pipe-delimited tables.
    """
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if _is_html_comment_line(line):
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            p = doc.add_heading("", level=len(m.group(1)))
            _add_inline_formatting(p, m.group(2).strip())
            i += 1
            continue

        if "|" in line and i + 1 < len(lines) and re.match(r"^[\s|:-]+$", lines[i + 1]):
            headers = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for ci, h in enumerate(headers):
                p = table.rows[0].cells[ci].paragraphs[0]
                p.text = ""
                _add_inline_formatting(p, h)
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci < len(headers):
                        p = table.rows[ri + 1].cells[ci].paragraphs[0]
                        p.text = ""
                        _add_inline_formatting(p, cell)
            continue

        m = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, m.group(1))
            i += 1
            continue

        m = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, m.group(1))
            i += 1
            continue

        m = re.match(r"^>\s?(.*)", line)
        if m:
            p = doc.add_paragraph(style="Intense Quote")
            _add_inline_formatting(p, m.group(1))
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        _add_inline_formatting(p, line)
        i += 1

    doc.save(path)
